"""
Processor module untuk membaca dan parse DOCX file.
Menampilkan isi dokumen secara lengkap dan terstruktur.
"""

import json
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, asdict
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

# ============ LOGGING SETUP ============
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class DocxElement:
    """Representasi element dalam DOCX"""
    type: str  # 'paragraph', 'table', 'image', 'text', etc
    content: str
    style: Optional[str] = None
    level: int = 0


@dataclass
class TableData:
    """Representasi data tabel"""
    rows: List[List[str]]
    columns: int
    rows_count: int


@dataclass
class DocumentContent:
    """Hasil parsing dokumen lengkap"""
    filename: str
    elements: List[Dict[str, Any]]
    total_paragraphs: int
    total_tables: int
    total_images: int
    text_content: str
    metadata: Dict[str, Any]


class DocxProcessor:
    """
    Processor untuk membaca dan parse file DOCX.
    
    Features:
    - Extract paragraf dengan formatting
    - Extract tabel dengan struktur lengkap
    - Extract gambar
    - Extract metadata
    - Support nested elements
    """
    
    def __init__(self, file_path: str):
        """
        Initialize processor dengan file path.
        
        Args:
            file_path: Path ke file DOCX
        """
        self.file_path = Path(file_path)
        self.document = None
        self.elements: List[Dict[str, Any]] = []
        self.images: List[Dict[str, Any]] = []
        self.tables: List[TableData] = []
        
        self._validate_file()
    
    def _validate_file(self) -> None:
        """Validasi file DOCX"""
        if not self.file_path.exists():
            raise FileNotFoundError(f"File tidak ditemukan: {self.file_path}")
        
        if self.file_path.suffix.lower() != '.docx':
            raise ValueError(f"File harus berformat DOCX, bukan {self.file_path.suffix}")
        
        logger.info(f"File valid: {self.file_path}")
    
    def process(self) -> DocumentContent:
        """
        Proses file DOCX dan return hasil lengkap.
        
        Returns:
            DocumentContent dengan semua isi dokumen
        """
        try:
            self.document = Document(self.file_path)
            logger.info("Dokumen berhasil dibuka")
            
            # Process semua element
            self._extract_elements()
            
            # Extract metadata
            metadata = self._extract_metadata()
            
            # Kumpulkan text content
            text_content = self._get_full_text()
            
            result = DocumentContent(
                filename=self.file_path.name,
                elements=self.elements,
                total_paragraphs=len([e for e in self.elements if e['type'] == 'paragraph']),
                total_tables=len(self.tables),
                total_images=len(self.images),
                text_content=text_content,
                metadata=metadata
            )
            
            logger.info(f"Processing selesai: {result.total_paragraphs} paragraf, "
                       f"{result.total_tables} tabel, {result.total_images} gambar")
            
            return result
        
        except Exception as e:
            logger.error(f"Error saat processing: {str(e)}", exc_info=True)
            raise
    
    def _extract_elements(self) -> None:
        """Extract semua elements dari dokumen"""
        for block in self.document.element.body:
            if isinstance(block, CT_P):
                self._process_paragraph(block)
            elif isinstance(block, CT_Tbl):
                self._process_table(block)
    
    def _process_paragraph(self, para_element: CT_P) -> None:
        """
        Process paragraph element.
        
        Args:
            para_element: CT_P element dari dokumen
        """
        try:
            para = Paragraph(para_element, self.document.element.body.getparent())
            
            # Extract text
            text = para.text.strip()
            
            if not text:
                return
            
            # Extract formatting
            style = para.style.name if para.style else "Normal"
            level = para.paragraph_format.outline_level or 0
            
            # Extract runs untuk detail formatting
            runs_info = []
            for run in para.runs:
                run_data = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size.pt if run.font.size else None,
                }
                runs_info.append(run_data)
            
            element_dict = {
                'type': 'paragraph',
                'content': text,
                'style': style,
                'level': level,
                'runs': runs_info,
                'alignment': str(para.alignment) if para.alignment else 'left',
                'space_before': para.paragraph_format.space_before,
                'space_after': para.paragraph_format.space_after,
            }
            
            # Extract images dari paragraph
            self._extract_images_from_paragraph(para, element_dict)
            
            self.elements.append(element_dict)
        
        except Exception as e:
            logger.warning(f"Error processing paragraph: {str(e)}")
    
    def _process_table(self, table_element: CT_Tbl) -> None:
        """
        Process table element.
        
        Args:
            table_element: CT_Tbl element dari dokumen
        """
        try:
            table = Table(table_element, self.document.element.body.getparent())
            
            rows_data = []
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for col_idx, cell in enumerate(row.cells):
                    cell_content = self._extract_cell_content(cell)
                    row_data.append(cell_content)
                rows_data.append(row_data)
            
            table_dict = {
                'type': 'table',
                'rows': rows_data,
                'rows_count': len(table.rows),
                'columns_count': len(table.columns),
                'content': self._table_to_string(rows_data),
            }
            
            self.elements.append(table_dict)
            self.tables.append(TableData(
                rows=rows_data,
                columns=len(table.columns),
                rows_count=len(table.rows)
            ))
        
        except Exception as e:
            logger.warning(f"Error processing table: {str(e)}")
    
    def _extract_cell_content(self, cell: _Cell) -> str:
        """
        Extract content dari cell.
        
        Args:
            cell: Cell dari tabel
            
        Returns:
            Text content dari cell
        """
        content = []
        for paragraph in cell.paragraphs:
            text = paragraph.text.strip()
            if text:
                content.append(text)
        
        return "\n".join(content) if content else ""
    
    def _table_to_string(self, rows_data: List[List[str]]) -> str:
        """
        Convert tabel ke string representation.
        
        Args:
            rows_data: Data baris tabel
            
        Returns:
            String representation tabel
        """
        if not rows_data:
            return ""
        
        # Hitung lebar kolom
        col_widths = []
        for col_idx in range(len(rows_data[0])):
            max_width = max(len(str(row[col_idx])) for row in rows_data)
            col_widths.append(max_width + 2)
        
        # Build string
        lines = []
        for row_idx, row in enumerate(rows_data):
            cells = []
            for col_idx, cell in enumerate(row):
                cell_str = str(cell).ljust(col_widths[col_idx])
                cells.append(cell_str)
            
            line = "|" + "|".join(cells) + "|"
            lines.append(line)
            
            # Add separator setelah header row
            if row_idx == 0:
                separator = "+" + "+".join(["-" * width for width in col_widths]) + "+"
                lines.insert(0, separator)
                lines.append(separator)
        
        lines.append(separator)
        return "\n".join(lines)
    
    def _extract_images_from_paragraph(self, para: Paragraph, element_dict: Dict) -> None:
        """
        Extract images dari paragraph.
        
        Args:
            para: Paragraph element
            element_dict: Dictionary untuk menyimpan info image
        """
        images_in_para = []
        
        for run in para.runs:
            for rel in run._element.getparent().iter():
                if 'drawing' in rel.tag or 'image' in rel.tag:
                    images_in_para.append({
                        'type': 'inline_image',
                        'tag': rel.tag
                    })
        
        if images_in_para:
            element_dict['images'] = images_in_para
            self.images.extend(images_in_para)
    
    def _extract_metadata(self) -> Dict[str, Any]:
        """
        Extract metadata dari dokumen.
        
        Returns:
            Dictionary berisi metadata
        """
        try:
            core_props = self.document.core_properties
            
            metadata = {
                'title': core_props.title or "N/A",
                'subject': core_props.subject or "N/A",
                'author': core_props.author or "N/A",
                'created': str(core_props.created) if core_props.created else "N/A",
                'modified': str(core_props.modified) if core_props.modified else "N/A",
                'comments': core_props.comments or "N/A",
                'keywords': core_props.keywords or "N/A",
            }
            
            return metadata
        
        except Exception as e:
            logger.warning(f"Error extracting metadata: {str(e)}")
            return {}
    
    def _get_full_text(self) -> str:
        """
        Ambil seluruh text content dari dokumen.
        
        Returns:
            Full text content
        """
        full_text = []
        
        for element in self.elements:
            if element['type'] == 'paragraph':
                full_text.append(element['content'])
            elif element['type'] == 'table':
                full_text.append(element['content'])
        
        return "\n\n".join(full_text)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert hasil processing ke dictionary"""
        result = self.process()
        return {
            'filename': result.filename,
            'elements': result.elements,
            'summary': {
                'total_paragraphs': result.total_paragraphs,
                'total_tables': result.total_tables,
                'total_images': result.total_images,
            },
            'metadata': result.metadata,
            'full_text': result.text_content,
        }
    
    def to_json(self, output_path: Optional[str] = None) -> str:
        """
        Export hasil ke JSON.
        
        Args:
            output_path: Path untuk menyimpan JSON (optional)
            
        Returns:
            JSON string
        """
        data = self.to_dict()
        json_str = json.dumps(data, indent=2, ensure_ascii=False)
        
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(json_str)
            logger.info(f"JSON exported ke: {output_path}")
        
        return json_str


def process_docx_file(file_path: str) -> DocumentContent:
    """
    Helper function untuk process file DOCX.
    
    Args:
        file_path: Path ke file DOCX
        
    Returns:
        DocumentContent dengan hasil parsing
    """
    processor = DocxProcessor(file_path)
    return processor.process()
